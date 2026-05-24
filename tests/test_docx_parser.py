from __future__ import annotations

from app.chunker import chunk_records
from app.parsers_docx import parse_docx


def test_docx_tables_keep_nearest_heading_context(tmp_path):
    from docx import Document

    doc_path = tmp_path / "policies.docx"
    doc = Document()
    doc.add_heading("Compensation Policy", level=1)
    doc.add_paragraph("Salary table:")
    salary = doc.add_table(rows=1, cols=3)
    salary.rows[0].cells[0].text = "Level"
    salary.rows[0].cells[1].text = "Base salary"
    salary.rows[0].cells[2].text = "Bonus"
    row = salary.add_row()
    row.cells[0].text = "Junior"
    row.cells[1].text = "15m"
    row.cells[2].text = "2m"

    doc.add_heading("Leave Policy", level=1)
    leave = doc.add_table(rows=1, cols=2)
    leave.rows[0].cells[0].text = "Year"
    leave.rows[0].cells[1].text = "Days"
    row = leave.add_row()
    row.cells[0].text = "1"
    row.cells[1].text = "12"
    doc.save(doc_path)

    records = parse_docx(doc_path)

    assert records[0]["text"] == "Compensation Policy\nSalary table:"
    assert records[0]["metadata"]["heading"] == "Compensation Policy"

    assert records[1]["text"] == (
        "Compensation Policy\nLevel: Junior | Base salary: 15m | Bonus: 2m"
    )
    assert records[1]["metadata"]["title"] == "Compensation Policy"
    assert records[1]["metadata"]["heading"] == "Compensation Policy"
    assert records[1]["metadata"]["table_idx"] == 1
    assert records[1]["metadata"]["row_num"] == 2

    assert records[2]["text"] == "Leave Policy\nYear: 1 | Days: 12"
    assert records[2]["metadata"]["title"] == "Leave Policy"
    assert records[2]["metadata"]["heading"] == "Leave Policy"
    assert records[2]["metadata"]["table_idx"] == 2
    assert records[2]["metadata"]["row_num"] == 2


def test_docx_heading_metadata_survives_chunking():
    chunks = chunk_records(
        [
            {
                "text": "Compensation Policy\nLevel: Junior | Base salary: 15m",
                "metadata": {
                    "title": "Compensation Policy",
                    "heading": "Compensation Policy",
                    "table_idx": 1,
                    "row_num": 2,
                    "lang": "en",
                },
            }
        ],
        kb_id=1,
        source_id="10",
        filename="policies.docx",
        file_type="docx",
        file_hash="hash",
        kb_version="v1",
        ingest_signature="sig",
    )

    assert chunks[0]["title"] == "Compensation Policy"
    assert chunks[0]["heading"] == "Compensation Policy"
    assert chunks[0]["table_idx"] == 1
    assert chunks[0]["row_num"] == 2
