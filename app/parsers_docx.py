"""
DOCX parser: extracts text section-by-section and table row-by-row in document order.
Requires: python-docx
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from app.lang import detect_language

logger = logging.getLogger(__name__)


def _normalize(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return "" if text.lower() in {"nan", "none", "null", ""} else text


def _iter_block_items(doc: Any):
    from docx.table import Table  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore

    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield Paragraph(child, doc)
        elif tag == "tbl":
            yield Table(child, doc)


def _is_heading(style_name: str) -> bool:
    return style_name.strip().lower().startswith("heading")


def _table_records(
    table: Any,
    *,
    table_idx: int,
    current_heading: str,
    fallback_title: str,
) -> list[dict[str, Any]]:
    if not table.rows:
        return []

    headers = [_normalize(cell.text) for cell in table.rows[0].cells]
    if not any(headers):
        return []

    title = current_heading or fallback_title
    records: list[dict[str, Any]] = []
    for row_idx, row in enumerate(table.rows[1:], start=2):
        values = [_normalize(cell.text) for cell in row.cells]
        parts = [
            f"{headers[i]}: {values[i]}"
            for i in range(min(len(headers), len(values)))
            if values[i]
        ]
        row_text = " | ".join(parts).strip()
        if not row_text:
            continue

        text = f"{current_heading}\n{row_text}".strip() if current_heading else row_text
        records.append(
            {
                "text": text,
                "metadata": {
                    "title": title,
                    "heading": current_heading or None,
                    "table_idx": table_idx,
                    "row_num": row_idx,
                    "lang": detect_language(text),
                },
            }
        )

    return records


def parse_docx(file_path: Path) -> list[dict[str, Any]]:
    try:
        from docx import Document  # type: ignore
    except ImportError as err:
        raise ValueError(
            "DOCX parsing requires python-docx. Install with: pip install python-docx"
        ) from err

    doc = Document(str(file_path))
    records: list[dict[str, Any]] = []
    current_heading = ""
    current_body: list[str] = []
    table_idx = 0

    def _flush_section() -> None:
        text = "\n".join(current_body).strip()
        if text:
            full_text = f"{current_heading}\n{text}".strip() if current_heading else text
            records.append(
                {
                    "text": full_text,
                    "metadata": {
                        "title": current_heading or file_path.stem,
                        "heading": current_heading or None,
                        "lang": detect_language(full_text),
                    },
                }
            )
        current_body.clear()

    for block in _iter_block_items(doc):
        if hasattr(block, "text"):
            style = block.style.name if block.style else ""
            text = block.text.strip()
            if not text:
                continue

            if _is_heading(style):
                _flush_section()
                current_heading = text
            else:
                current_body.append(text)
            continue

        _flush_section()
        table_idx += 1
        records.extend(
            _table_records(
                block,
                table_idx=table_idx,
                current_heading=current_heading,
                fallback_title=file_path.stem,
            )
        )

    _flush_section()

    logger.info(
        "Parsed DOCX: %s records (%s paragraphs, %s tables) from %s",
        len(records),
        len(doc.paragraphs),
        len(doc.tables),
        file_path.name,
    )
    return records
